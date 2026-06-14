#!/usr/bin/env python3
"""Apply SCI v3 text and statistics updates to the current final thesis DOCX."""

from __future__ import annotations

import csv
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "00_论文最终材料_2026-06-13" / "01_最终论文"
CURRENT_V3 = OUT_DIR / "硕士论文_SCI_v3优化版.docx"
FALLBACK_SRC = ROOT / "00_论文最终材料_2026-06-12" / "01_最终论文" / "硕士论文_最终检查版.docx"
OUT = OUT_DIR / "硕士论文_SCI_v3实验完成回填版.docx"
STAT_CSV = ROOT / "SCI一区升级规划" / "统计检验_v3" / "complete_stat_tests_current.csv"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def add_rich_text(paragraph, text: str, size_pt: float | None = None, bold: bool | None = None) -> None:
    i = 0
    while i < len(text):
        if i + 1 < len(text) and text[i] in {"I", "G"} and text[i + 1] in {"1", "2"}:
            run = paragraph.add_run(text[i])
            set_run_font(run, size_pt, bold)
            sub = paragraph.add_run(text[i + 1])
            set_run_font(sub, size_pt, bold)
            sub.font.subscript = True
            i += 2
            continue
        if i + 1 < len(text) and text[i : i + 2] == "R0":
            run = paragraph.add_run("R")
            set_run_font(run, size_pt, bold)
            sub = paragraph.add_run("0")
            set_run_font(sub, size_pt, bold)
            sub.font.subscript = True
            i += 2
            continue
        run = paragraph.add_run(text[i])
        set_run_font(run, size_pt, bold)
        i += 1


def set_paragraph_text(paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    add_rich_text(paragraph, text)


def find_para(doc: Document, exact: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact:
            return paragraph
    raise RuntimeError(f"paragraph not found: {exact}")


def insert_paragraph_after(anchor, text: str, style: str = "Normal"):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = anchor._parent.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    new_para = anchor._parent._insert_paragraph_before(None)
    # python-docx has no stable insert-after API. Replace the temporary node
    # with a normal paragraph object by moving its XML after the anchor.
    anchor._p.addnext(new_para._p)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        add_rich_text(new_para, text)
    return new_para


def insert_after(anchor, element) -> None:
    anchor._p.addnext(element)


def make_paragraph(doc: Document, text: str, style: str = "Normal", align=None):
    paragraph = doc.add_paragraph(style=style)
    clear_paragraph(paragraph)
    add_rich_text(paragraph, text)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def move_paragraph_after(paragraph, anchor) -> None:
    anchor._p.addnext(paragraph._p)


def move_table_after(table, anchor) -> None:
    anchor._p.addnext(table._tbl)


def selected_stat_rows() -> list[dict[str, str]]:
    rows = read_csv(STAT_CSV)
    wanted = [
        "T=15 双轨辟谣干预",
        "T=30 双轨辟谣干预",
        "T=60 双轨辟谣干预",
        "线上单维辟谣干预",
        "线下单维辟谣干预",
        "弱双轨辟谣干预强度",
        "强双轨辟谣干预强度",
    ]
    result = []
    for name in wanted:
        for row in rows:
            if row["scenarioName"] == name and row["metric"] == "cumulativeInfectedPersonHours_right":
                result.append(row)
                break
    return result


def add_stat_test_table(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("表 5.9") and "统计检验结果" in paragraph.text:
            return
    anchor = find_para(doc, "5.3.4 参数敏感性与消融实验")
    insertion_anchor = anchor

    caption_zh = make_paragraph(doc, "表 5.9  主实验累计传播者人时的统计检验结果", align=WD_ALIGN_PARAGRAPH.CENTER)
    caption_en = make_paragraph(
        doc,
        "Table 5.9  Statistical Tests for Cumulative Spreader-Person-Hours in Main Experiments",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    fields = ["场景", "均值", "降幅/%", "Welch t", "BH-FDR p", "Hedges g"]
    stat_rows = selected_stat_rows()
    table = doc.add_table(rows=1, cols=len(fields))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, field in enumerate(fields):
        cell = table.rows[0].cells[idx]
        set_paragraph_text(cell.paragraphs[0], field)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 9, True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in stat_rows:
        values = [
            row["scenarioName"],
            row["mean"],
            row["reductionPctVsBaseline"],
            row["welchT"],
            row["pBHByMetric"],
            row["hedgesG"],
        ]
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_paragraph_text(cells[idx].paragraphs[0], value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cells[idx].paragraphs[0].runs:
                set_run_font(run, 8.5, False)

    note = make_paragraph(
        doc,
        "表 5.9 进一步给出累计传播者人时的统计检验结果。T=15、T=30 和 T=60 双轨辟谣干预的累计负荷降幅均通过多重比较修正后的显著性检验，但效应大小不同；其中 T=15 的降幅和效应量最大，说明早期干预对传播持续负担的压缩更明显。需要注意的是，这一统计结论只针对本文仿真参数和随机种子集合成立，不能直接外推为现实政策因果效果。",
    )

    # Move in reverse order so final order is caption_zh, caption_en, table, note, then old anchor.
    move_paragraph_after(note, insertion_anchor)
    move_table_after(table, insertion_anchor)
    move_paragraph_after(caption_en, insertion_anchor)
    move_paragraph_after(caption_zh, insertion_anchor)


def replace_target_paragraphs(doc: Document) -> None:
    replacements = {
        "为了避免某一次随机仿真结果影响判断，本文对主实验场景都进行了30个随机种子的重复运行，并把无干预基准情景作为比较参照。比较指标包括传播峰值、累计传播者人时和效应量。其中，累计传播者人时反映的是传播者数量与持续时间叠加后的总体负担。结果显示，T=15双轨辟谣干预的累计传播者人时均值为23567.50，较无干预基准情景下降42.82%；T=30和T=60的累计负荷降幅分别为24.64%和6.04%。这说明，干预越早启动，越容易影响传播者高峰的形成；如果干预启动较晚，它更多影响的是高峰之后传播还会持续多久。":
            "为了避免某一次随机仿真结果影响判断，本文对主实验场景进行了 30 个随机种子的重复运行，并把无干预基准情景作为比较参照。比较指标包括传播峰值、累计传播者人时、Welch t 检验、p 值和效应量。其中，累计传播者人时反映的是传播者数量与持续时间叠加后的总体负担。结果显示，T=15 双轨辟谣干预的累计传播者人时均值为 23567.50，较无干预基准情景下降 42.82%；T=30 和 T=60 的累计负荷降幅分别为 24.64% 和 6.04%。这说明，干预越早启动，越容易影响传播者高峰的形成；如果干预启动较晚，它更多影响的是高峰之后传播还会持续多久。",
        "从效应量看，T=15 双轨辟谣干预、强双轨辟谣干预强度和 T=30 双轨辟谣干预对累计传播者人时的影响更明显。换言之，这些情景不仅让某一个时间点的传播者数量发生变化，也压缩了传播持续的总负担。相比之下，T=30、T=60 和单维干预对最大峰值的降低较有限。这一结果说明，当传播高峰已经接近形成时，干预未必能明显压低最高点，但仍可能缩短传播者保持活跃的时间，从而降低累计影响。":
            "从效应量看，T=15 双轨辟谣干预、强双轨辟谣干预强度和 T=30 双轨辟谣干预对累计传播者人时的影响更明显。换言之，这些情景不仅让某一个时间点的传播者数量发生变化，也压缩了传播持续的总负担。相比之下，T=30、T=60 和单维干预对最大峰值的降低较有限，且部分峰值指标在 Welch 检验中并不显著。这一结果说明，当传播高峰已经接近形成时，干预未必能明显压低最高点，但仍可能缩短传播者保持活跃的时间，从而降低累计影响。",
        "参数设定仍会影响仿真结果的解释范围。第五章已经补充 30 seed 重复实验、参数敏感性和消融比较，并使用 CHECKED 数据进行了线上传播代理曲线的初步校准。这些工作提高了结果的稳定性和外部参照性，但仍不能把模型参数直接等同于现实传播参数。尤其是 I2、G2 等线下传播者或线下辟谣干预状态，目前缺少可以直接观测的数据，其解释仍然依赖模型假设和代理变量。":
            "参数设定仍会影响仿真结果的解释范围。第五章已经补充 30 seed 重复实验、统计检验、参数敏感性和消融比较，并使用 CHECKED 数据进行了线上传播代理曲线的初步校准。这些工作提高了结果的稳定性和外部参照性，但仍不能把模型参数直接等同于现实传播参数。尤其是 I2、G2 等线下传播者或线下辟谣干预状态，目前缺少可以直接观测的数据，其解释仍然依赖模型假设和代理变量。若按 SCI 一区标准继续推进，还需要完成传统 SEIR、单层网络 SEIR、无 G 状态/无耦合模型三类基线的同口径比较。",
        "未来研究还可在仿真层面拓展情景比较。本文已完成一组基础敏感性和消融实验，后续可以进一步采用拉丁超立方抽样、Sobol 指数或其他全局敏感性方法，对更大的参数范围进行系统评估。这样做的目的，是判断哪些参数真正主导峰值、达峰时间和累计传播负荷，而不是只比较少量预设情景 [73][74]。如果能够把这些方法与真实数据校准结合起来，模型结论的稳健性会更强。":
            "未来研究还可在仿真层面拓展情景比较。本文已完成一组基础敏感性和消融实验，并已经整理 Morris 筛选和拉丁超立方抽样的全局敏感性实验队列。后续可以在完整模型和三类基线模型上继续运行这些队列，系统评估传播速率、干预启动时刻、退出时间、平均接触度和网络重连概率等参数对峰值、达峰时间和累计传播负荷的影响。这样做的目的，是判断哪些参数真正主导结论，而不是只比较少量预设情景。",
        "未来研究可进一步结合真实案例开展外部验证。在具备数据条件和伦理合规前提下，可选择具有清晰首发时间、明确辟谣节点、可追踪传播链和官方信息发布时间的公开事件，比较模型输出与真实传播过程之间的对应关系。此类研究应继续避免把模型结果简单等同于现实因果证明，而应关注模型能否合理解释传播趋势、干预时机和传播空间差异。":
            "未来研究可进一步结合真实案例开展外部验证。在具备数据条件和伦理合规前提下，可选择具有清晰首发时间、明确辟谣节点、可追踪传播链和官方信息发布时间的公开事件，比较模型输出与真实传播过程之间的对应关系。除 CHECKED 外，还可以选取带有发布时间和谣言标签的公开 COVID-19 谣言数据集，或按突发事件组织的 PHEME 谣言数据集，作为不重新大幅调参的外部趋势验证。此类研究应继续避免把模型结果简单等同于现实因果证明，而应关注模型能否合理解释传播趋势、干预时机和传播空间差异。",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text])


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    src = CURRENT_V3 if CURRENT_V3.exists() else FALLBACK_SRC
    shutil.copy2(src, OUT)
    doc = Document(OUT)
    replace_target_paragraphs(doc)
    add_stat_test_table(doc)
    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
