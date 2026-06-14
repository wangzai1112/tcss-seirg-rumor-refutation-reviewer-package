#!/usr/bin/env python3
"""Insert baseline-model comparison results into the SCI v3 thesis DOCX."""

from __future__ import annotations

import csv
import shutil
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "00_论文最终材料_2026-06-13" / "01_最终论文"
SRC = OUT_DIR / "硕士论文_SCI_v3实验完成回填版.docx"
FALLBACK_SRC = OUT_DIR / "硕士论文_SCI_v3优化版.docx"
OUT = OUT_DIR / "硕士论文_SCI_v3完整基线回填版.docx"
STAT_CSV = ROOT / "SCI一区升级规划" / "统计检验_v3" / "complete_stat_tests_current.csv"

MODEL_NAMES = {
    "0": "耦合 SEIRG",
    "1": "传统 SEIR",
    "2": "单层网络 SEIR",
    "3": "无 G/无耦合",
}

SCENARIO_LABELS = {
    "no_intervention": "无干预",
    "T15": "T=15 干预",
    "T30": "T=30 干预",
    "T60": "T=60 干预",
}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def add_rich_text(paragraph, text: str, size_pt: float | None = None, bold: bool | None = None) -> None:
    i = 0
    while i < len(text):
        if i + 1 < len(text) and text[i] in {"I", "G"} and text[i + 1] in {"1", "2"}:
            run = paragraph.add_run(text[i])
            set_run_font(run, size_pt, bold)
            sub = paragraph.add_run(text[i + 1])
            set_run_font(sub, size_pt, bold)
            sub.font.subscript = True
            i += 2
            continue
        run = paragraph.add_run(text[i])
        set_run_font(run, size_pt, bold)
        i += 1


def set_paragraph_text(paragraph, text: str, size_pt: float | None = None, bold: bool | None = None) -> None:
    clear_paragraph(paragraph)
    add_rich_text(paragraph, text, size_pt, bold)


def make_paragraph(doc: Document, text: str, style: str = "Normal", align=None, bold: bool | None = None):
    paragraph = doc.add_paragraph(style=style)
    clear_paragraph(paragraph)
    add_rich_text(paragraph, text, bold=bold)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def move_paragraph_before(paragraph, anchor) -> None:
    anchor._p.addprevious(paragraph._p)


def move_table_before(table, anchor) -> None:
    anchor._p.addprevious(table._tbl)


def find_anchor(doc: Document):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "5.4 本章小结":
            return paragraph
    raise RuntimeError("anchor not found: 5.4 本章小结")


def section_exists(doc: Document) -> bool:
    return any("5.3.6 基线模型对比" in paragraph.text for paragraph in doc.paragraphs)


def scenario_key(name: str) -> str:
    if "无干预" in name or "no_intervention" in name:
        return "no_intervention"
    if "T=15" in name or "T15" in name:
        return "T15"
    if "T=30" in name or "T30" in name:
        return "T30"
    if "T=60" in name or "T60" in name:
        return "T60"
    return name


def build_rows() -> list[dict[str, str]]:
    metrics: dict[tuple[str, str, str], dict[str, str]] = {}
    for row in read_csv(STAT_CSV):
        mode = row.get("modelMode", "0") or "0"
        key = scenario_key(row.get("scenarioName", ""))
        metric = row.get("metric", "")
        if key not in SCENARIO_LABELS:
            continue
        metrics[(mode, key, metric)] = row

    table_rows: list[dict[str, str]] = []
    for mode in ["0", "1", "2", "3"]:
        for key in ["no_intervention", "T15", "T30", "T60"]:
            peak = metrics.get((mode, key, "totalInfectedPeak"))
            burden = metrics.get((mode, key, "cumulativeInfectedPersonHours_right"))
            if not peak or not burden:
                continue
            table_rows.append(
                {
                    "模型": MODEL_NAMES.get(mode, f"M{mode}"),
                    "情景": SCENARIO_LABELS[key],
                    "n": burden.get("n", ""),
                    "传播峰值均值": peak.get("mean", ""),
                    "累计传播者人时": burden.get("mean", ""),
                    "累计负荷降幅/%": burden.get("reductionPctVsBaseline", ""),
                    "BH-FDR p": burden.get("pBHByMetric", ""),
                    "效应量 g": burden.get("hedgesG", ""),
                }
            )
    return table_rows


def summarize_rows(rows: list[dict[str, str]]) -> str:
    by_model: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        by_model[row["模型"]].append(row)
    parts: list[str] = []
    for model, model_rows in by_model.items():
        intervention_rows = [row for row in model_rows if row["情景"] != "无干预"]
        if not intervention_rows:
            continue
        best = max(
            intervention_rows,
            key=lambda row: float(row["累计负荷降幅/%"] or "-999"),
        )
        parts.append(f"{model} 中累计负荷降幅最大的情景为 {best['情景']}，降幅为 {best['累计负荷降幅/%']}%。")
    if not parts:
        return "三类基线模型尚未形成完整统计结果，后续应在相同随机种子和相同指标口径下继续补齐。"
    return " ".join(parts)


def add_baseline_section(doc: Document) -> None:
    if section_exists(doc):
        return
    anchor = find_anchor(doc)
    rows = build_rows()

    section_title = make_paragraph(doc, "5.3.6 基线模型对比实验", style="Heading 3")
    intro = make_paragraph(
        doc,
        "为判断本文耦合 SEIRG 模型的结果是否只是由个别结构设定造成，本文进一步设置三类基线模型进行同口径比较：传统 SEIR 模型、单层网络 SEIR 模型，以及去除 G 状态和层间耦合后的简化模型。三类基线均采用与主模型一致的随机种子、干预启动时刻和统计指标，重点比较传播峰值、累计传播者人时、显著性检验和效应量。",
    )
    caption_zh = make_paragraph(doc, "表 5.10  不同模型结构下的基线对比统计结果", align=WD_ALIGN_PARAGRAPH.CENTER)
    caption_en = make_paragraph(
        doc,
        "Table 5.10  Baseline Comparison Results across Alternative Model Structures",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    fields = ["模型", "情景", "n", "传播峰值均值", "累计传播者人时", "累计负荷降幅/%", "BH-FDR p", "效应量 g"]
    table = doc.add_table(rows=1, cols=len(fields))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, field in enumerate(fields):
        cell = table.rows[0].cells[idx]
        set_paragraph_text(cell.paragraphs[0], field, 8.5, True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, field in enumerate(fields):
            set_paragraph_text(cells[idx].paragraphs[0], row.get(field, ""), 8, False)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    explanation = make_paragraph(
        doc,
        "表 5.10 的作用不是说明某一种模型一定最接近现实，而是检验本文结论是否依赖耦合网络、G 状态和层间转移等结构设定。若主模型与三类基线在累计传播负荷、效应量和显著性方向上保持一致，说明早期干预降低传播持续负担的结论具有一定结构稳健性；若差异较大，则需要回到模型机制层面解释差异来源。"
        + summarize_rows(rows),
    )

    for element in [section_title, intro, caption_zh, caption_en]:
        move_paragraph_before(element, anchor)
    move_table_before(table, anchor)
    move_paragraph_before(explanation, anchor)


def main() -> None:
    source = SRC if SRC.exists() else FALLBACK_SRC
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, OUT)
    doc = Document(OUT)
    add_baseline_section(doc)
    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
